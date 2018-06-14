import csv
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

inFile = open('tmp.txt')
inReader = csv.reader(inFile)
id_pwds = list(inReader)

doc = docx.Document()

style = doc.styles['Normal']
font = style.font
font.name = 'MSゴシック'
font.size = Pt(24)

listStyle = doc.styles['List']
listFont = listStyle.font
listFont.size = Pt(34)

footStyle = doc.styles['List 2']
footFont = footStyle.font
footFont.size = Pt(20)

section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE

new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

for (id, pwd) in id_pwds:
    doc.add_heading('Linux実技試験用　アカウント', 0)
    doc.add_paragraph('以下のユーザー名とパスワードを用いて、Linuxにログインし、実技問題を解答すること。')
    doc.add_paragraph(' ')

    p1 = doc.add_paragraph('ユーザー名      ')
    p1.add_run(id).bold = True
    p1.style = listStyle

    p2 = doc.add_paragraph('パスワード      ')
    p2.add_run(pwd).bold = True
    p2.style = listStyle

    doc.add_paragraph('')

    pn = doc.add_paragraph()
    pn.add_run('学籍番号                              氏名                                      .').underline = True
    pf = doc.add_paragraph('退室する際は、この用紙を、教員が指定した場所に提出すること。')
    pf.style = footStyle

    doc.add_page_break()

doc.save('id_pass.docx')
